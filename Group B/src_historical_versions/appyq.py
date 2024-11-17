import dash
from dash import dcc, html, Input, Output, State
import pandas as pd
import plotly.express as px
import io
import dash_bootstrap_components as dbc
from dash import dash_table
import folium
from dash_extensions import BeforeAfter
import dash_leaflet as dl
import pdfkit
import docx

# Initialize the app
app = dash.Dash(__name__, external_stylesheets=[dbc.themes.MINTY], suppress_callback_exceptions=True)
server = app.server

# Define app layout
app.layout = html.Div([
    dcc.Location(id='url', refresh=False),
    html.Div(id='page-content')
])

# Define layout for Page 1 - Data Upload/Input
page_1_layout = dbc.Container([
    dbc.NavbarSimple(
        brand="GHG Emissions Calculator",
        brand_href="/",
        color="primary",
        dark=True,
        className="mb-5"
    ),
    dbc.Card([
        dbc.CardHeader(html.H2("Data Upload")),
        dbc.CardBody([
            dbc.Row([
                dbc.Col([
                    dcc.Upload(
                        id='upload-data',
                        children=html.Div([
                            'Drag and Drop or ',
                            html.A('Select Files', style={'color': '#007bff', 'textDecoration': 'underline'})
                        ]),
                        style={
                            'width': '100%', 'height': '60px', 'lineHeight': '60px',
                            'borderWidth': '2px', 'borderStyle': 'dashed',
                            'borderRadius': '5px', 'textAlign': 'center', 'margin': '10px'
                        },
                        multiple=False
                    ),
                    html.Div(id='file-upload-status', style={'marginTop': '10px'}),
                ], width=6),
            ]),
            dbc.Row([
                dbc.Col([
                    html.H5("Building Name"),
                    dcc.Input(id='building-input', type='text', placeholder="Enter Building Name", style={'width': '100%'}),
                ], width=4),
                dbc.Col([
                    html.H5("Year"),
                    dcc.Input(id='year-input', type='number', placeholder="Enter Year", style={'width': '100%'}),
                ], width=4),
            ], className="mb-4"),
            dbc.Row([
                dbc.Col([
                    html.H5("Latitude"),
                    dcc.Input(id='latitude-input', type='number', placeholder="Enter Latitude", style={'width': '100%'}),
                ], width=4),
                dbc.Col([
                    html.H5("Longitude"),
                    dcc.Input(id='longitude-input', type='number', placeholder="Enter Longitude", style={'width': '100%'}),
                ], width=4),
            ], className="mb-4"),
            dbc.Row([
                dbc.Col([
                    html.H5("Address"),
                    dcc.Input(id='address-input', type='text', placeholder="Enter Address", style={'width': '100%'}),
                ], width=8),
            ], className="mb-4"),
            dbc.Row([
                dbc.Col([
                    html.H5("Type of Function (Building)"),
                    dcc.Input(id='type-function-input', type='text', placeholder="Enter Building Function Type", style={'width': '100%'}),
                ], width=4),
                dbc.Col([
                    html.H5("Size"),
                    dcc.Dropdown(
                        id='size-input',
                        options=[
                            {'label': 'Large', 'value': 'Large'},
                            {'label': 'Medium', 'value': 'Medium'},
                            {'label': 'Small', 'value': 'Small'}
                        ],
                        placeholder="Select Size",
                        style={'width': '100%'}
                    ),
                ], width=4),
            ], className="mb-4"),
            dbc.Row([
                dbc.Col([
                    html.H5("Number of Employees"),
                    dcc.Input(id='employee-input', type='number', placeholder="Enter Number of Employees", style={'width': '100%'}),
                ], width=4),
                dbc.Col([
                    html.H5("Transportation (CO2 Emissions in kg)"),
                    dcc.Input(id='transportation-input', type='number', placeholder="Enter Transportation CO2 Emissions", style={'width': '100%'}),
                ], width=4),
            ], className="mb-4"),
            dbc.Row([
                dbc.Col([
                    html.H5("Gross Floor Area (GFA in m²)"),
                    dcc.Input(id='gfa-input', type='number', placeholder="Enter GFA", style={'width': '100%'}),
                ], width=4),
                dbc.Col([
                    html.H5("Energy (in kWh)"),
                    dcc.Input(id='energy-input', type='number', placeholder="Enter Energy Consumption", style={'width': '100%'}),
                ], width=4),
            ], className="mb-4"),
            dbc.Row([
                dbc.Col([
                    html.H5("Waste (in kg)"),
                    dcc.Input(id='waste-input', type='number', placeholder="Enter Waste Generated", style={'width': '100%'}),
                ], width=4),
                dbc.Col([
                    html.H5("Water (in m³)"),
                    dcc.Input(id='water-input', type='number', placeholder="Enter Water Consumption", style={'width': '100%'}),
                ], width=4),
            ], className="mb-4"),
            dbc.Row([
                dbc.Col([
                    html.H5("Scope 1 Emissions"),
                    dcc.Input(id='scope1-input', type='number', placeholder="Enter Scope 1 emissions", style={'width': '100%'}),
                    dbc.Tooltip("Direct emissions from owned or controlled sources.", target='scope1-input')
                ], width=4),
                dbc.Col([
                    html.H5("Scope 2 Emissions"),
                    dcc.Input(id='scope2-input', type='number', placeholder="Enter Scope 2 emissions", style={'width': '100%'}),
                    dbc.Tooltip("Indirect emissions from the generation of purchased electricity.", target='scope2-input')
                ], width=4),
                dbc.Col([
                    html.H5("Scope 3 Emissions"),
                    dcc.Input(id='scope3-input', type='number', placeholder="Enter Scope 3 emissions", style={'width': '100%'}),
                    dbc.Tooltip("All other indirect emissions that occur in the value chain.", target='scope3-input')
                ], width=4),
            ], className="mb-4"),
            dbc.Row([
                dbc.Col([
                    dbc.Button("Next: Data Quality Check", id="next-button", color="success", className="mr-2", href='/page-2')
                ], width=2),
            ]),
        ])
    ], className="mb-5")
])

# Define layout for Page 2 - Data Quality Check and EDA
page_2_layout = dbc.Container([
    dbc.NavbarSimple(
        brand="Data Quality Check and EDA",
        brand_href="/page-2",
        color="primary",
        dark=True,
        className="mb-5"
    ),
    dbc.Row([
        dbc.Col(html.H1("Data Quality Check and Exploratory Data Analysis", className="text-center"), className="mb-5 mt-5")
    ]),
    dbc.Row([
        dbc.Col(html.Div(id="eda-output", className="mt-4")),
    ]),
    dbc.Row([
        dbc.Col([
            html.H5("Visualization of Data Completeness"),
            dcc.Graph(id='completeness-graph'),
        ]),
    ], className="mt-5"),
    dbc.Row([
        dbc.Col([
            html.H5("Data Table Overview"),
            dash_table.DataTable(id='data-table', style_table={'overflowX': 'auto'})
        ])
    ], className="mt-4"),
    dbc.Row([
        dbc.Col([
            dbc.Button("Re-upload Data", id="reupload-button", color="warning", href='/page-1'),
            dbc.Button("Next: Model Running and Comparison", id="next-model-button", color="primary", href='/page-3', className="ml-2")
        ], width=6),
    ])
])

# Define layout for Page 3 - Model Running and Visualization
# Dummy benchmarking data (initial benchmarks)

page_3_layout = dbc.Container([
    dbc.NavbarSimple(
        brand="Model Results and Benchmarking",
        brand_href="/page-3",
        color="primary",
        dark=True,
        className="mb-5"
    ),
    dbc.Row([
        dbc.Col(html.H1("Model Results and Benchmarking", className="text-center"), className="mb-5 mt-5")
    ]),
    dbc.Row([
        dbc.Col([
            html.H5("Model Results Visualization"),
            dcc.Graph(id='model-results-graph'),
        ]),
    ], className="mt-5"),
    dbc.Row([
        dbc.Col([
            html.H5("Scope 1 Emissions Breakdown"),
            dcc.Graph(id='scope1-pie-chart')
        ], width=4),
        dbc.Col([
            html.H5("Scope 2 Emissions Breakdown"),
            dcc.Graph(id='scope2-pie-chart')
        ], width=4),
        dbc.Col([
            html.H5("Scope 3 Emissions Breakdown"),
            dcc.Graph(id='scope3-pie-chart')
        ], width=4),
    ], className="mt-5"),
    dbc.Row([
        dbc.Col([
            html.H5("Benchmarking Visualization"),
            dcc.Graph(id='benchmarking-graph', figure=px.bar(pd.DataFrame({
                'Building': ['Building A', 'Building B', 'Building C'],
                'Emissions Intensity': [120, 150, 110]
            }), x='Building', y='Emissions Intensity', title="Emissions Intensity Benchmarking"))
        ])
    ], className="mt-5"),
    dbc.Row([
        dbc.Col([
            dbc.Button("Next: Generate Report", id="next-report-button", color="primary", href='/page-4')
        ], width=4),
    ])
])

# Define layout for Page 4 - Report Generation and Chatbot
page_4_layout = dbc.Container([
    dbc.NavbarSimple(
        brand="Generate Report and Chat Assistance",
        brand_href="/page-4",
        color="primary",
        dark=True,
        className="mb-5"
    ),
    dbc.Row([
        dbc.Col(html.H1("Generate Report and Chat Assistance", className="text-center"), className="mb-5 mt-5")
    ]),
    dbc.Row([
        dbc.Col([
            html.H5("Generate Report"),
            html.H5("Select Report Format", className="text-center"),
            dcc.RadioItems(
                id='report-format',
                options=[
                    {'label': 'PDF', 'value': 'pdf'},
                    {'label': 'Word', 'value': 'docx'},
                    {'label': 'HTML', 'value': 'html'}
                ],
                value='pdf',  # Default option
                inline=True,
                labelStyle={'padding-right': '15px'},
                style={'display': 'block', 'textAlign': 'center'}
            ),
            dbc.Button("Download Report", id="report-button", color="secondary", className="mt-3"),
            dcc.Download(id="download-report")
        ], width=4),
    ], className="mb-5"),
    dbc.Row([
        dbc.Col([
            html.H5("Chatbot Assistance"),
            dcc.Input(id='chatbot-input', type='text', placeholder="Ask a question about the report...", style={'width': '100%'}),
            html.Div(id='chatbot-response', className="mt-3")
        ], width=6),
    ])
])

# Define layout for Page 5 - Geospatial Analysis
page_5_layout = dbc.Container([
    dbc.NavbarSimple(
        brand="Geospatial Analysis of Buildings",
        brand_href="/page-5",
        color="primary",
        dark=True,
        className="mb-5"
    ),
    dbc.Row([
        dbc.Col(html.H1("Geospatial Analysis of Buildings", className="text-center"), className="mb-5 mt-5")
    ]),
    dbc.Row([
        dbc.Col([
            html.H5("Building Locations on Map"),
            dl.Map(center=[1.3521, 103.8198], zoom=12, children=[
                dl.TileLayer(),
                dl.LayerGroup(id="layer-group")
            ], style={'width': '100%', 'height': '500px', 'margin': "auto", "display": "block"})
        ])
    ], className="mt-5"),
    dbc.Row([
        dbc.Col([
            dbc.Button("Back to Data Upload", id="back-upload-button", color="primary", href='/page-1')
        ], width=4),
    ])
])

# Update page content based on URL
@app.callback(Output('page-content', 'children'),
              Input('url', 'pathname'))
def display_page(pathname):
    if pathname == '/page-2':
        return page_2_layout
    elif pathname == '/page-3':
        return page_3_layout
    elif pathname == '/page-4':
        return page_4_layout
    elif pathname == '/page-5':
        return page_5_layout
    else:
        return page_1_layout


# generate report

@app.callback(
    Output('download-report', 'data'),
    [Input('report-button', 'n_clicks')],
    [State('report-format', 'value')]
)
def generate_report(n_clicks, report_format):
    if n_clicks:
        try:
            # Generate a blank HTML report
            if report_format == 'html':
                report_html = """
                <html>
                    <head><title>Blank Report</title></head>
                    <body><h1>Blank Report</h1><p>This is a placeholder for the report.</p></body>
                </html>
                """
                with open('report.html', 'w') as f:
                    f.write(report_html)
                return dcc.send_file('report.html')

            # Generate a blank PDF report
            elif report_format == 'pdf':
                report_html = """
                <html>
                    <head><title>Blank PDF Report</title></head>
                    <body><h1>Blank PDF Report</h1><p>This is a placeholder for the PDF report.</p></body>
                </html>
                """
                pdfkit.from_string(report_html, 'report.pdf')
                return dcc.send_file('report.pdf')

            # Generate a blank Word report
            elif report_format == 'docx':
                doc = docx.Document()
                doc.add_heading('Blank Word Report', 0)
                doc.add_paragraph("This is a placeholder for the Word report.")
                doc.save('report.docx')
                return dcc.send_file('report.docx')

        except Exception as e:
            print(f"Error during report generation: {e}")
            return None



# Data upload callback
@app.callback(
    Output('file-upload-status', 'children'),
    Input('upload-data', 'contents'),
    State('upload-data', 'filename')
)
def handle_file_upload(contents, filename):
    if contents:
        # Process the uploaded file here
        return f"Uploaded file: {filename}"
    else:
        return "No file uploaded yet."

# Additional callbacks for EDA, model results, benchmarking, geospatial visualization, and report generation can be added similarly

# Run the app
if __name__ == '__main__':
    app.run_server(debug=True)